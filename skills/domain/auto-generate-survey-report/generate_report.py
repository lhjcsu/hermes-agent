#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
============================================================
  威海高区保税物流中心（B型）项目
  岩土工程勘察报告 — 自动生成工具（完整版v2）
============================================================
从华宁勘察软件导出的 Excel 数据自动填充到报告模板中，
一键生成完整的岩土工程勘察报告（.docx）。

用法:
    cd "D:\项目\威海高区保税物流中心（B型）项目"
    python generate_report.py

输出:
    D:\项目\威海高区保税物流中心（B型）项目\25019威海高区保税物流中心_正式报告.docx
============================================================
"""

import os
import re
import sys
import xlrd
import openpyxl
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict

# ============================================================
# 0. 路径配置
# ============================================================
BASE_DIR = r"D:\项目\威海高区保税物流中心（B型）项目"
EXCEL_DIR = os.path.join(BASE_DIR, r"已有资料\excel表格")
TEMPLATE_PATH = os.path.join(BASE_DIR, "25019威海高区保税物流中心.docx")
OUTPUT_PATH = os.path.join(BASE_DIR, "25019威海高区保税物流中心_正式报告.docx")


# ============================================================
# 1. 数据读取函数
# ============================================================

def read_xls(filepath, sheet_idx=0):
    """读取 .XLS 老格式文件，返回 (sheet_name, rows_list)"""
    wb = xlrd.open_workbook(filepath)
    sname = wb.sheet_names()[sheet_idx]
    ws = wb.sheet_by_name(sname)
    rows = []
    for r in range(ws.nrows):
        row = [ws.cell(r, c).value for c in range(ws.ncols)]
        rows.append(row)
    return sname, rows


def read_xlsx(filepath, sheet_name=None):
    """读取 .xlsx 文件，返回 (sheet_name, rows_list)"""
    wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
    if sheet_name:
        sname = sheet_name
    else:
        sname = wb.sheetnames[0]
    ws = wb[sname]
    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append(list(row))
    wb.close()
    return sname, rows


def safe_float(v, default=None):
    if v is None:
        return default
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if s == '' or s == '-' or s == '—':
        return default
    try:
        return float(s)
    except:
        return default


def safe_str(v):
    if v is None:
        return ''
    return str(v).strip()


def format_val(v, fmt='.2f'):
    if v is None:
        return ''
    try:
        return f'{float(v):{fmt}}'
    except:
        return str(v)


def format_val_int(v):
    """格式化为整数"""
    if v is None:
        return ''
    try:
        fv = float(v)
        if fv == int(fv):
            return str(int(fv))
        return f'{fv:.1f}'
    except:
        return str(v)


# ============================================================
# 2. 数据解析器
# ============================================================

class ProjectData:
    """解析并存储所有华宁 Excel 数据"""

    def __init__(self):
        self.project_name = "威海高区保税物流中心（B型）项目"
        self.boreholes = []
        self.borehole_type_stats = {}
        self.workload_summary = {}
        self.layer_stats = {}
        self.spt_raw_data = []
        self.spt_stats = {}
        self.cpt_raw_data = []
        self.cpt_stats = {}
        self.physical_stats_data = {}
        self.liquefaction_data = []
        self.rock_data = []

    def load_all(self):
        print("正在读取所有 Excel 数据...")
        self._load_boreholes()
        self._load_workload_stats()
        self._load_spt()
        self._load_cpt()
        self._load_physical_stats()
        self._load_liquefaction()
        self._load_rock()
        self._compute_borehole_types()
        print("所有数据读取完成。")

    def _load_boreholes(self):
        path = os.path.join(EXCEL_DIR, "勘探点一览表19.XLS")
        _, rows = read_xls(path)
        self.boreholes = []
        for row in rows[6:]:
            if len(row) < 14:
                continue
            seq = safe_str(row[0])
            if not seq:
                continue
            if seq == '119' or seq == '合计':
                break
            bh = {
                'seq': seq, 'id': safe_str(row[1]), 'type': safe_str(row[2]),
                'elevation': safe_float(row[3]), 'depth': safe_float(row[4]),
                'water_table_depth': safe_float(row[7]), 'water_table_elv': safe_float(row[8]),
                'undisturbed_samples': safe_float(row[9], 0),
                'disturbed_samples': safe_float(row[10], 0),
                'rock_samples': safe_float(row[11], 0),
                'spt_count': safe_float(row[12], 0), 'n63_5_count': safe_float(row[13], 0),
            }
            self.boreholes.append(bh)
        print(f"  - 勘探点: {len(self.boreholes)} 个")

    def _compute_borehole_types(self):
        type_count = defaultdict(int)
        for bh in self.boreholes:
            type_count[bh['type']] += 1
        self.borehole_type_stats = dict(type_count)

        self.workload_summary = {
            'total_boreholes': len(self.boreholes),
            'total_depth': sum(bh['depth'] or 0 for bh in self.boreholes),
            'borehole_types': type_count,
            'undisturbed_samples': int(sum(bh['undisturbed_samples'] or 0 for bh in self.boreholes)),
            'disturbed_samples': int(sum(bh['disturbed_samples'] or 0 for bh in self.boreholes)),
            'rock_samples': int(sum(bh['rock_samples'] or 0 for bh in self.boreholes)),
            'spt_tests': int(sum(bh['spt_count'] or 0 for bh in self.boreholes)),
            'n63_tests': int(sum(bh['n63_5_count'] or 0 for bh in self.boreholes)),
        }

        self.标贯孔 = type_count.get('标贯孔', 0)
        self.取土孔 = type_count.get('取土孔', 0)
        self.一般性钻孔 = type_count.get('一般性钻孔', 0)
        self.重探孔 = sum(v for k, v in type_count.items() if '重探' in k or '动力' in k)

        self.标贯孔进尺 = sum(bh['depth'] or 0 for bh in self.boreholes if bh['type'] == '标贯孔')
        self.取土孔进尺 = sum(bh['depth'] or 0 for bh in self.boreholes if bh['type'] == '取土孔')
        self.一般钻孔进尺 = sum(bh['depth'] or 0 for bh in self.boreholes if bh['type'] == '一般性钻孔')
        self.重探孔进尺 = sum(bh['depth'] or 0 for bh in self.boreholes if '重探' in bh['type'] or '动力' in bh['type'])

        total = self.workload_summary['total_depth']
        print(f"  - 钻孔类型: {dict(type_count)}, 总进尺: {total:.2f}m")

    def _load_workload_stats(self):
        path = os.path.join(EXCEL_DIR, "场地各孔厚度层底深度标高及层顶深度标高统计明细表1.xlsx")
        try:
            _, t1 = read_xlsx(path, '表1')
            _, t2 = read_xlsx(path, '表2')
            _, t3 = read_xlsx(path, '表3')

            layer_ids = []
            for c in range(1, len(t1[2])):
                v = safe_str(t1[2][c])
                if v:
                    layer_ids.append(v)

            for idx, lid in enumerate(layer_ids, 1):
                stats = {}
                if len(t1) > 4:
                    stats['thickness_min'] = safe_float(t1[4][idx]) if idx < len(t1[4]) else None
                    stats['thickness_max'] = safe_float(t1[5][idx]) if len(t1) > 5 and idx < len(t1[5]) else None
                    stats['thickness_n'] = safe_float(t1[6][idx]) if len(t1) > 6 and idx < len(t1[6]) else None
                    stats['thickness_avg'] = safe_float(t1[7][idx]) if len(t1) > 7 and idx < len(t1[7]) else None
                if len(t2) > 4:
                    stats['depth_min'] = safe_float(t2[4][idx]) if idx < len(t2[4]) else None
                    stats['depth_max'] = safe_float(t2[5][idx]) if len(t2) > 5 and idx < len(t2[5]) else None
                    stats['depth_avg'] = safe_float(t2[7][idx]) if len(t2) > 7 and idx < len(t2[7]) else None
                if len(t3) > 4:
                    stats['elv_min'] = safe_float(t3[4][idx]) if idx < len(t3[4]) else None
                    stats['elv_max'] = safe_float(t3[5][idx]) if len(t3) > 5 and idx < len(t3[5]) else None
                    stats['elv_avg'] = safe_float(t3[7][idx]) if len(t3) > 7 and idx < len(t3[7]) else None
                self.layer_stats[lid] = stats
            print(f"  - 地层统计: {len(self.layer_stats)} 层")
        except Exception as e:
            print(f"  ! 地层统计读取异常: {e}")

    def _load_spt(self):
        path = os.path.join(EXCEL_DIR, "标准贯入试验成果统计表19.XLS")
        try:
            _, rows = read_xls(path)
            self.spt_raw_data = rows[6:]
            self.spt_stats = {}
            i = 6
            while i < len(rows):
                row = rows[i]
                layer = safe_str(row[0])
                if not layer:
                    i += 1
                    continue
                stat_check = safe_str(row[1])
                if stat_check in ('最小值', '最 小 值'):
                    lid = layer
                    self.spt_stats[lid] = {}
                    # 实测值
                    self.spt_stats[lid]['min_raw'] = safe_float(row[6])
                    self.spt_stats[lid]['min_corr'] = safe_float(row[7]) if len(row) > 7 else None
                    if i + 1 < len(rows):
                        self.spt_stats[lid]['max_raw'] = safe_float(rows[i+1][6])
                        self.spt_stats[lid]['max_corr'] = safe_float(rows[i+1][7]) if len(rows[i+1]) > 7 else None
                    if i + 2 < len(rows):
                        self.spt_stats[lid]['n'] = safe_float(rows[i+2][6])
                    if i + 3 < len(rows):
                        self.spt_stats[lid]['avg_raw'] = safe_float(rows[i+3][6])
                        self.spt_stats[lid]['avg_corr'] = safe_float(rows[i+3][7]) if len(rows[i+3]) > 7 else None
                    if i + 4 < len(rows):
                        self.spt_stats[lid]['std'] = safe_float(rows[i+4][6])
                    if i + 5 < len(rows):
                        self.spt_stats[lid]['cv'] = safe_float(rows[i+5][6])
                    if i + 6 < len(rows):
                        self.spt_stats[lid]['std_val_raw'] = safe_float(rows[i+6][6])
                        self.spt_stats[lid]['std_val_corr'] = safe_float(rows[i+6][7]) if len(rows[i+6]) > 7 else None
                    i += 7
                else:
                    i += 1
            print(f"  - 标贯数据: {len(self.spt_raw_data)} 条, {len(self.spt_stats)} 层统计")
        except Exception as e:
            print(f"  ! 标贯数据读取异常: {e}")

    def _load_cpt(self):
        path = os.path.join(EXCEL_DIR, "动力触探N63.5试验成果统计表19.XLS.XLS")
        try:
            _, rows = read_xls(path)
            self.cpt_raw_data = rows[5:]
            self.cpt_stats = {}
            for i, row in enumerate(rows[5:]):
                layer = safe_str(row[0])
                if not layer:
                    continue
                if safe_str(row[1]) == '最小值':
                    lid = layer
                    self.cpt_stats[lid] = {}
                    if i < len(rows) - 5:
                        self.cpt_stats[lid]['n'] = safe_float(rows[5+i+2][5])
                        self.cpt_stats[lid]['max_raw'] = safe_float(rows[5+i+1][5])
                        self.cpt_stats[lid]['min_raw'] = safe_float(row[5])
                        self.cpt_stats[lid]['avg_raw'] = safe_float(rows[5+i+3][5])
                        self.cpt_stats[lid]['std'] = safe_float(rows[5+i+4][5])
                        self.cpt_stats[lid]['cv'] = safe_float(rows[5+i+5][5])
                        self.cpt_stats[lid]['std_val_raw'] = safe_float(rows[5+i+6][5])
            print(f"  - 动探数据: {len(self.cpt_raw_data)} 条, {len(self.cpt_stats)} 层统计")
        except Exception as e:
            print(f"  ! 动探数据读取异常: {e}")

    def _load_physical_stats(self):
        """读取物理力学性质指标统计表19.XLS"""
        path = os.path.join(EXCEL_DIR, "物理力学性质指标统计表19.XLS")
        try:
            _, rows = read_xls(path)

            stat_names = {
                '最 小 值': 'min', '最 大 值': 'max', '数据个数': 'n',
                '平均值': 'avg', '标 准 差': 'std', '变异系数': 'cv', '标 准 值': 'std_val',
            }

            # 基于第6行列头的列->指标名映射
            # 对于粘土层(3,4,6,8)有完整数据:
            # cols 3-5: depth/elv/thickness (跳过)
            # col 6: W(%), col 7: Gs, col 8: γ, col 9: γd, col 10: e0, col 11: Sr(%),
            # col 12: WL(%), col 13: WP(%), col 14: IP, col 15: IL,
            # col 16: w/wL, col 17: Ir, col 18: wL/e, col 19: c(kPa)
            # 实际上，第8行(row 8)开始的列顺序：
            # 对杂填土(层1): 只有前3列有数据(depth/elv/thickness)
            # 对细砂(层2): 只有前3列
            # 对淤泥质粉质黏土(层3): 从col3(depth)到col19(c)
            col_to_indicator = {
                3: 'depth(m)', 4: 'elv(m)', 5: 'thickness(m)',
                6: 'W(%)', 7: 'Gs', 8: 'γ(kN/m3)', 9: 'γd(kN/m3)',
                10: 'e0', 11: 'Sr(%)',
                12: 'WL(%)', 13: 'WP(%)', 14: 'IP', 15: 'IL',
                16: 'w/wL', 17: 'Ir', 18: 'wL/e', 19: 'c(kPa)',
            }

            current_layer = None
            for row in rows[8:]:
                lid = safe_str(row[0])
                stat_type = safe_str(row[2])
                if lid:
                    current_layer = lid
                stat_key = stat_names.get(stat_type)
                if not stat_key or not current_layer:
                    continue
                if current_layer not in self.physical_stats_data:
                    self.physical_stats_data[current_layer] = {}

                for c in range(3, min(len(row), 20)):
                    v = safe_float(row[c])
                    if v is not None:
                        indicator = col_to_indicator.get(c, f'col{c}')
                        if indicator not in self.physical_stats_data[current_layer]:
                            self.physical_stats_data[current_layer][indicator] = {}
                        self.physical_stats_data[current_layer][indicator][stat_key] = v

            print(f"  - 物理力学统计: {len(self.physical_stats_data)} 层")
        except Exception as e:
            print(f"  ! 物理力学统计读取异常: {e}")

    def _load_liquefaction(self):
        path = os.path.join(EXCEL_DIR, "标准贯入试验液化判别及液化指数计算成果表_19.XLS")
        try:
            _, rows = read_xls(path)
            data_rows = []
            for row in rows[10:]:
                bh = safe_str(row[0])
                if bh:
                    data_rows.append(row)
            self.liquefaction_data = data_rows
            print(f"  - 液化判别: {len(self.liquefaction_data)} 条")
        except Exception as e:
            print(f"  ! 液化判别读取异常: {e}")

    def _load_rock(self):
        path = os.path.join(EXCEL_DIR, "岩石试验指标分层统计表19.XLS")
        try:
            _, rows = read_xls(path)
            data_rows = []
            for row in rows[7:]:
                layer = safe_str(row[0])
                if not layer:
                    continue
                stat_keywords = ['最小值', '最大值', '数据个数', '平均值', '标准差', '变异系数', '标准值']
                if any(k in safe_str(row[1]) for k in stat_keywords):
                    break
                data_rows.append(row)
            self.rock_data = data_rows
            print(f"  - 岩石试验: {len(self.rock_data)} 条")
        except Exception as e:
            print(f"  ! 岩石试验读取异常: {e}")

    def get_layer_stat(self, layer_num, field='thickness'):
        layer_map = {
            '1': '1', '2': '2', '3': '3', '4': '4', '5': '5',
            '6': '6', '7': '7', '8': '8', '9': '9',
            '10': '10-1', '11': '10-2', '12': '10-3',
        }
        lid = layer_map.get(str(layer_num), str(layer_num))
        stats = self.layer_stats.get(lid, {})
        prefix = field + '_'
        return {
            'min': stats.get(prefix + 'min', ''),
            'max': stats.get(prefix + 'max', ''),
            'avg': stats.get(prefix + 'avg', ''),
            'n': stats.get(prefix + 'n', ''),
        }

    def get_water_table_stats(self):
        depths = [bh['water_table_depth'] for bh in self.boreholes if bh['water_table_depth']]
        elvs = [bh['water_table_elv'] for bh in self.boreholes if bh['water_table_elv']]
        if not depths:
            return {'n': 0, 'depth_min': 0, 'depth_max': 0, 'depth_avg': 0,
                    'elv_min': 0, 'elv_max': 0, 'elv_avg': 0}
        return {
            'n': len(depths),
            'depth_min': min(depths), 'depth_max': max(depths),
            'depth_avg': sum(depths) / len(depths),
            'elv_min': min(elvs) if elvs else 0, 'elv_max': max(elvs) if elvs else 0,
            'elv_avg': sum(elvs) / len(elvs) if elvs else 0,
        }


# ============================================================
# 3. 报告填充器
# ============================================================

class ReportFiller:
    def __init__(self, data: ProjectData):
        self.data = data
        print(f"  正在打开模板: {TEMPLATE_PATH}")
        self.doc = Document(TEMPLATE_PATH)

    def fill(self):
        print("正在填充报告模板...")
        self._fill_paragraphs()
        self._fill_tables()
        print("报告填充完成。")

    def _set_cell(self, table, row_idx, col_idx, value):
        """安全设置单元格文本"""
        if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
            table.rows[row_idx].cells[col_idx].text = str(value)

    def _set_paragraph_text(self, p, text):
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = text
        else:
            p.clear()
            p.add_run(text)

    def _fill_paragraphs(self):
        paragraphs = self.doc.paragraphs
        d = self.data

        # 地层描述 - 填充统计数据
        layer_descs = {
            135: (1, '杂填土',
                '黄褐色、灰褐色，松散、局部中密，强度不均，主要成分风化岩渣土、碎石及建筑垃圾，块石在各孔分布不均，块石含量20~30%，径多为10~30cm，最大80cm左右，局部深度以块石、建筑垃圾为主。回填时间8年，尚未完成自重固结。'),
            139: (2, '细砂',
                '灰黑色、灰色，顶部松散，中下部稍密，饱和，主要成分长石、石英及少量云母，夹少量贝壳碎片，颗粒较均匀。'),
            142: (3, '淤泥质粉质黏土',
                '灰～灰黑色，软塑，土质较均匀，局部夹粉质黏土和细砂薄层，含贝壳碎片，稍具腥臭味。'),
            147: (4, '粉质黏土',
                '黄灰色~灰褐色，可塑，粉粒及黏粒组成，土质较均匀，局部夹粉细砂透镜体薄层。无摇振反应，切面光滑，干强度及韧性中等。'),
            150: (5, '中细砂',
                '浅灰色、黄褐色，中密，饱和，主要成分长石、石英，颗粒不均，含少量砾粒，局部夹粉质黏土薄层，磨圆度较差。'),
            154: (6, '粉质黏土',
                '黄褐色，硬塑，粉粒及黏粒组成，土质不均匀，中下部含少量中粗砂，局部夹细砂透镜体薄层。无摇振反应，切面光滑，干强度及韧性中等。'),
            157: (7, '中粗砂',
                '灰褐色、灰白色，密实，饱和。主要成分长石、石英，颗粒不均，夹粗砾砂，含砾石，粒径2～5mm左右，磨圆度较差。'),
            161: (8, '粉质黏土',
                '浅黄褐色、棕褐色，硬塑~坚硬，以黏性土为主，含较多砂粒及砾石，无摇振反应，切面有光泽，干强度、韧性高。'),
            164: (9, '中粗砂',
                '黄褐色、棕褐色，密实，饱和，主要成分为长石、石英，颗粒不均匀，磨圆度差，底部夹大量的砾石、碎石及风化碎屑。'),
        }

        for idx, (layer_num, name, base_desc) in layer_descs.items():
            if idx < len(paragraphs):
                t = d.get_layer_stat(layer_num, 'thickness')
                dep = d.get_layer_stat(layer_num, 'depth')
                elv = d.get_layer_stat(layer_num, 'elevation')

                text = f"{layer_num}层{name}（{base_desc}），"
                if t['min'] != '':
                    text += f"厚度:{format_val(t['min'])}～{format_val(t['max'])}m,平均{format_val(t['avg'])}m;"
                if dep['min'] != '':
                    text += f"层底埋深:{format_val(dep['min'])}～{format_val(dep['max'])}m,平均{format_val(dep['avg'])}m;"
                if elv['min'] != '':
                    text += f"层底标高:{format_val(elv['min'])}～{format_val(elv['max'])}m,平均{format_val(elv['avg'])}m。"
                # Remove trailing semicolon if ends with it
                if text.endswith(';'):
                    text = text[:-1] + '。'
                elif not text.endswith('。'):
                    text += '。'

                # 检查原段落是否有其他内容需要保留
                orig = paragraphs[idx].text.strip()
                if '场区' in orig and '场区' not in text:
                    # 保留"场区普遍分布"等信息
                    pass

                self._set_paragraph_text(paragraphs[idx], text)

        # 基岩层描述 (保留原格式, 仅更新内部数字)
        # 这些段落不替换，因为包含非结构化文本

        # 更新工作量段落
        if 104 < len(paragraphs):
            txt = (f"本次勘察实际完成钻孔{len(d.boreholes)}个，"
                   f"其中控制孔45个，一般孔{len(d.boreholes)-45}个。"
                   f"其中取土孔{d.取土孔}个、标贯孔{d.标贯孔}个，鉴别孔{d.一般性钻孔}个，波速孔6个勘察工作量统计表见表3-1。")
            self._set_paragraph_text(paragraphs[104], txt)

        # 更新地下水段落
        ws = d.get_water_table_stats()
        if 184 < len(paragraphs):
            txt = (f"勘察期间测得钻孔内水位埋深约{ws['depth_min']:.2f}~{ws['depth_max']:.2f}m，"
                   f"水位标高{ws['elv_min']:.2f}~{ws['elv_max']:.2f}m，详见表5-13。"
                   f"根据场地地下水的埋藏条件，地下水类型为第四系孔隙潜水，"
                   f"主要含水层为第1层素填土、第2层中细砂和第5层中细砂、第7层中粗砂、第9层中粗砂。")
            self._set_paragraph_text(paragraphs[184], txt)

        # 更新液化段落
        if 208 < len(paragraphs):
            liquefied = 0
            total = 0
            li_vals = []
            for row in d.liquefaction_data:
                if len(row) > 9:
                    layer = safe_str(row[1])
                    result = safe_str(row[9])
                    if layer == '2':
                        total += 1
                        if '液' in result:
                            liquefied += 1
                        if len(row) > 15:
                            li = safe_float(row[15])
                            if li and li > 0:
                                li_vals.append(li)
            if total > 0:
                li_min = min(li_vals) if li_vals else 0
                li_max = max(li_vals) if li_vals else 0
                txt = (f"第2层中细砂进行液化判别{total}个，液化{liquefied}个，不液化{total-liquefied}个，"
                       f"液化指数ILE为{li_min:.2f}~{li_max:.2f}，液化等级为轻微；"
                       f"第5层中细砂进行液化判别34个点，第7层中粗砂进行液化判别8个点，第9层中粗砂1个点，均不液化。"
                       f"液化判别详见标准贯入试验液化判别及液化指数计算成果表。")
                self._set_paragraph_text(paragraphs[208], txt)

        print("  - 段落更新完成")

    def _fill_tables(self):
        tables = self.doc.tables

        # 表3: 勘察工作量统计表 (16行x4列)
        if len(tables) > 3:
            self._fill_workload_table(tables[3])

        # 表4: 动探统计-杂填土 (3行x8列)
        if len(tables) > 4:
            self._fill_cpt_table(tables[4], '1')

        # 表5: 标贯-细砂 (5行x8列)
        if len(tables) > 5:
            self._fill_spt_table(tables[5], '2')

        # 表6: 物理力学-淤泥质粉质黏土 (13行x9列)
        if len(tables) > 6:
            self._fill_phys_table(tables[6], '3')

        # 表7: 物理力学-粉质黏土(4层)
        if len(tables) > 7:
            self._fill_phys_table(tables[7], '4')

        # 表8: 标贯-中细砂(5层) (5行x8列)
        if len(tables) > 8:
            self._fill_spt_table(tables[8], '5')

        # 表9: 物理力学-粉质黏土(6层)
        if len(tables) > 9:
            self._fill_phys_table(tables[9], '6')

        # 表10: 标贯-中粗砂(7层)
        if len(tables) > 10:
            self._fill_spt_table(tables[10], '7')

        # 表11: 物理力学-粉质黏土(8层)
        if len(tables) > 11:
            self._fill_phys_table(tables[11], '8')

        # 表12: 标贯-中粗砂(9层)
        if len(tables) > 12:
            self._fill_spt_table(tables[12], '9')

        # 表13: 标贯-全风化片麻岩(10-1) (3行x8列)
        if len(tables) > 13:
            self._fill_spt_rock_table(tables[13], '10-1')

        # 表14: 标贯-强风化片麻岩(10-2)
        if len(tables) > 14:
            self._fill_spt_rock_table(tables[14], '10-2')

        # 表15: 点荷载-碎块状强风化片麻岩(10-3) - 保留原数据

        # 表16: 地下水位统计表
        if len(tables) > 16:
            self._fill_water_table(tables[16])

        # 表18: 承载力建议值表 (15行x9列)
        if len(tables) > 18:
            self._fill_bearing_table(tables[18])

        print("  - 表格填充完成")

    def _fill_workload_table(self, table):
        d = self.data
        ws = d.workload_summary

        # Row索引和数据列索引(第4列=col 3)
        data_map = {
            1: str(len(d.boreholes)),           # 测定勘探点
            2: f"{ws['total_depth']:.2f}/{len(d.boreholes)}",  # 钻探 m/孔
            3: str(d.标贯孔),                    # 标贯孔
            4: str(d.取土孔),                    # 取土孔
            5: str(d.重探孔),                    # 动力触探孔
            6: str(d.一般性钻孔),                # 一般性钻孔
            7: str(ws['n63_tests']),            # 动探米数
            8: str(ws['spt_tests']),            # 标贯次数
            9: str(ws['undisturbed_samples']),  # 原状样
            10: str(ws['disturbed_samples']),   # 扰动样
            11: str(ws['rock_samples']),        # 岩样
            12: str(ws['undisturbed_samples'] + ws['disturbed_samples']),  # 土样数
        }

        for row_idx, val in data_map.items():
            if row_idx < len(table.rows) and len(table.rows[row_idx].cells) >= 4:
                table.rows[row_idx].cells[3].text = val

    def _fill_spt_table(self, table, layer_id):
        """填充标贯统计表 (5行x8列): 行0=header, 行1=header2, 行2=实测值, 行3=Es, 行4=fak"""
        stats = self.data.spt_stats.get(layer_id, {})
        if not stats:
            return

        # 行2: 标贯实测值
        # 列: 1=最小值, 2=最大值, 3=数据个数, 4=平均值, 5=标准差, 6=变异系数, 7=标准值
        self._set_cell(table, 2, 1, format_val(stats.get('min_raw')))
        self._set_cell(table, 2, 2, format_val(stats.get('max_raw')))
        self._set_cell(table, 2, 3, format_val_int(stats.get('n')))
        self._set_cell(table, 2, 4, format_val(stats.get('avg_raw')))
        self._set_cell(table, 2, 5, format_val(stats.get('std')))
        self._set_cell(table, 2, 6, format_val(stats.get('cv')))
        self._set_cell(table, 2, 7, format_val(stats.get('std_val_raw')))

        # 修正值行（如果有）
        if stats.get('avg_corr') is not None:
            self._set_cell(table, 3, 1, format_val(stats.get('min_corr')))
            self._set_cell(table, 3, 2, format_val(stats.get('max_corr')))
            self._set_cell(table, 3, 4, format_val(stats.get('avg_corr')))
            self._set_cell(table, 3, 7, format_val(stats.get('std_val_corr')))

        # Es和fak行
        bearing_defaults = {
            '2': ('10.0', '130'), '3': ('3.5', '70'), '4': ('7.0', '180'),
            '5': ('15.0', '200'), '6': ('8.0', '210'), '7': ('20.0', '280'),
            '8': ('9.0', '240'), '9': ('25.0', '350'),
        }
        es, fak = bearing_defaults.get(layer_id, ('', ''))
        self._set_cell(table, 3, 4, es)
        self._set_cell(table, 4, 4, fak)

    def _fill_spt_rock_table(self, table, layer_id):
        """填充基岩标贯表 (3行x8列)"""
        stats = self.data.spt_stats.get(layer_id, {})
        if not stats:
            return

        # Row 1: 标贯实测值
        self._set_cell(table, 2, 1, format_val(stats.get('min_raw')))
        self._set_cell(table, 2, 2, format_val(stats.get('max_raw')))
        self._set_cell(table, 2, 3, format_val_int(stats.get('n')))
        self._set_cell(table, 2, 4, format_val(stats.get('avg_raw')))
        self._set_cell(table, 2, 5, format_val(stats.get('std')))
        self._set_cell(table, 2, 6, format_val(stats.get('cv')))
        self._set_cell(table, 2, 7, format_val(stats.get('std_val_raw')))

    def _fill_cpt_table(self, table, layer_id):
        """填充动探统计表 (3行x8列) - 列顺序: 统计个数,最大值,最小值,平均值,标准差,变异系数,标准值"""
        stats = self.data.cpt_stats.get(layer_id, {})
        if not stats:
            return

        # Row 2: 数据行 (重型动力触探修正值)
        # 列: 1=统计个数, 2=最大值, 3=最小值, 4=平均值, 5=标准差, 6=变异系数, 7=标准值
        self._set_cell(table, 2, 1, format_val_int(stats.get('n')))
        self._set_cell(table, 2, 2, format_val(stats.get('max_raw')))
        self._set_cell(table, 2, 3, format_val(stats.get('min_raw')))
        self._set_cell(table, 2, 4, format_val(stats.get('avg_raw')))
        self._set_cell(table, 2, 5, format_val(stats.get('std')))
        self._set_cell(table, 2, 6, format_val(stats.get('cv')))
        self._set_cell(table, 2, 7, format_val(stats.get('std_val_raw')))

    def _fill_phys_table(self, table, layer_id):
        """填充物理力学指标统计表 (13行x9列)"""
        data = self.data.physical_stats_data.get(layer_id, {})
        if not data:
            return

        # 指标映射: (行索引, 数据key在physical_stats中的名称)
        # 列: 2=最小值, 3=最大值, 4=平均值, 5=n, 6=std, 7=cv, 8=std_val
        row_map = {
            1: 'W(%)',       # 含水率
            2: 'γ(kN/m3)',   # 重度
            3: 'e0',         # 孔隙比
            4: 'WL(%)',      # 液限
            5: 'WP(%)',      # 塑限
            6: 'IP',         # 塑性指数
            7: 'IL',         # 液性指数
            8: 'c(kPa)',     # 黏聚力
        }

        for row_idx, indicator in row_map.items():
            if indicator not in data:
                continue
            idata = data[indicator]
            if row_idx >= len(table.rows):
                continue
            self._set_cell(table, row_idx, 2, format_val(idata.get('min')))
            self._set_cell(table, row_idx, 3, format_val(idata.get('max')))
            self._set_cell(table, row_idx, 4, format_val(idata.get('avg')))
            self._set_cell(table, row_idx, 5, format_val_int(idata.get('n')))
            self._set_cell(table, row_idx, 6, format_val(idata.get('std')))
            self._set_cell(table, row_idx, 7, format_val(idata.get('cv')))
            self._set_cell(table, row_idx, 8, format_val(idata.get('std_val')))

    def _fill_water_table(self, table):
        ws = self.data.get_water_table_stats()
        if len(table.rows) > 1:
            row = table.rows[1]
            cells = row.cells
            if len(cells) >= 7:
                cells[0].text = str(ws['n'])
                cells[1].text = f"{ws['depth_min']:.2f}"
                cells[2].text = f"{ws['depth_max']:.2f}"
                cells[3].text = f"{ws['depth_avg']:.2f}"
                cells[4].text = f"{ws['elv_min']:.2f}"
                cells[5].text = f"{ws['elv_max']:.2f}"
                cells[6].text = f"{ws['elv_avg']:.2f}"

    def _fill_bearing_table(self, table):
        bearing_data = [
            (1, "杂填土", "/", "/", "/", "/", "19.0", "3.0", "18"),
            (2, "细砂", "130", "10.0", "", "", "19.5", "", "28"),
            (3, "淤泥质粉质黏土", "70", "3.5", "", "", "17.5", "15", "8"),
            (4, "粉质黏土", "180", "7.0", "", "", "19.2", "20", "15"),
            (5, "中细砂", "200", "15.0", "", "", "19.5", "", "30"),
            (6, "粉质黏土", "210", "8.0", "", "", "19.5", "36", "18"),
            (7, "中粗砂", "280", "20.0", "", "", "20.0", "", "32"),
            (8, "粉质黏土", "240", "9.0", "", "", "19.8", "32", "20"),
            (9, "中粗砂", "350", "25.0", "", "", "20.5", "", "35"),
            ("10-1", "全风化片麻岩", "450", "40.0", "", "", "22.0", "", "35"),
            ("10-2", "强风化片麻岩", "700", "60.0", "", "", "23.0", "", "40"),
            ("10-3", "碎块状强风化片麻岩", "1200", "", "", "", "24.0", "", "45"),
        ]

        for i, (layer_id, name, fak, es1, es, e0, gamma, c, phi) in enumerate(bearing_data):
            row_idx = i + 2
            if row_idx < len(table.rows):
                cells = table.rows[row_idx].cells
                if len(cells) >= 9:
                    cells[0].text = str(layer_id)
                    cells[1].text = name
                    cells[2].text = fak
                    cells[3].text = es1
                    cells[4].text = es
                    cells[5].text = e0
                    cells[6].text = gamma
                    cells[7].text = c
                    cells[8].text = phi


# ============================================================
# 4. 主函数
# ============================================================

def main():
    print("=" * 60)
    print("  岩土工程勘察报告自动生成工具（完整版v2）")
    print("  项目: 威海高区保税物流中心（B型）项目")
    print("=" * 60)
    print()

    data = ProjectData()
    data.load_all()

    filler = ReportFiller(data)
    filler.fill()

    print(f"\n正在保存报告到: {OUTPUT_PATH}")
    filler.doc.save(OUTPUT_PATH)
    print(f"报告已生成: {OUTPUT_PATH}")
    print()
    print("生成完成!")


if __name__ == '__main__':
    main()
